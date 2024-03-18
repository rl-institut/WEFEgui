from datetime import date
from docx import Document
from docx.table import Table
from docx.shape import InlineShape
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement, ns
from docx.opc.constants import RELATIONSHIP_TYPE
import pandas as pd
import numpy as np
import numpy_financial as npf
import os
import json
import csv
import base64
import io
from cp_nigeria.models import ConsumerGroup, DemandTimeseries, Options
from projects.models import Asset
from projects.constants import ENERGY_DENSITY_DIESEL, CURRENCY_SYMBOLS
from business_model.models import EquityData, BusinessModel, BMAnswer
from business_model.helpers import B_MODELS
from dashboard.models import FancyResults, KPIScalarResults
from projects.models import EconomicData
from django.shortcuts import get_object_or_404
from django.db.models import Func, Sum
from django.contrib.staticfiles.storage import staticfiles_storage
from django.templatetags.static import static
from dashboard.models import get_costs


class Unnest(Func):
    subquery = True  # contains_subquery = True for Django > 4.2.7
    function = "unnest"


HOUSEHOLD_TIERS = [
    ("", "Remove SHS threshold"),
    ("very_low", "Very Low Consumption Estimate"),
    ("low", "Low Consumption Estimate"),
    ("middle", "Middle Consumption Estimate"),
    ("high", "High Consumption Estimate"),
    ("very_high", "Very High Consumption Estimate"),
]


def csv_to_dict(filepath, label_col="label"):
    # the csv must contain a column named "label" containing the variable name, which will be used to construct the
    # nested dictionaries
    dict = {}
    if os.path.exists(staticfiles_storage.path(filepath)) is True:
        with open(staticfiles_storage.path(filepath), encoding="utf-8") as csvfile:
            csvreader = csv.reader(csvfile, delimiter=",", quotechar='"')
            for i, row in enumerate(csvreader):
                if i == 0:
                    hdr = row
                    label_idx = hdr.index(label_col)
                else:
                    label = row[label_idx]
                    dict[label] = {}
                    for k, v in zip(hdr, row):
                        dict[label][k] = v

    return dict


FINANCIAL_PARAMS = csv_to_dict("financial_tool/financial_parameters_list.csv")
OUTPUT_PARAMS = csv_to_dict("cpn_output_params.csv")


def set_table_format(param, from_dict=OUTPUT_PARAMS):
    param_dict = {}
    for item in ["verbose", "description", "unit"]:
        if param == "":
            param_dict[item] = ""
            continue
        if item == "description":
            dict_value = "" if from_dict[param][item] == "" else help_icon(from_dict[param][item])
        else:
            dict_value = from_dict[param][item]

        if "currency" in dict_value:
            # TODO make this custom depending on project currency
            dict_value = dict_value.replace("currency", CURRENCY_SYMBOLS["NGN"])

        param_dict[item] = dict_value

    return param_dict


def get_project_summary(project):
    bm_name = BusinessModel.objects.get(scenario=project.scenario).model_name
    options = Options.objects.get(project=project)
    ft = FinancialTool(project)
    if "inverter" in ft.system_params["supply_source"].tolist():
        inverter_aggregated_flow = ft.system_params.loc[
            (ft.system_params["category"] == "total_flow") & (ft.system_params["supply_source"] == "inverter"), "value"
        ].iloc[0]
    else:
        inverter_aggregated_flow = 0

    yearly_production = ft.yearly_production_electricity
    total_investments = ft.total_capex("NGN")
    total_demand, peak_demand, daily_demand = get_demand_indicators(project)
    renewable_share = inverter_aggregated_flow / total_demand * 100
    project_lifetime = project.economic_data.duration
    bm_name = B_MODELS[bm_name]["Verbose"]
    tariff = EquityData.objects.get(scenario=project.scenario).estimated_tariff * ft.exchange_rate
    if options.community is not None:
        community_name = options.community.name
    else:
        community_name = project.name

    currency_symbol = project.economic_data.currency_symbol

    project_summary = {
        "project_name": f"{project.name}",
        # "Community name": f"{community_name}",
        "location": f"{project.latitude:.4f}°, {project.longitude:.4f}°",
        "yearly_production": f"{yearly_production:,.0f} kWh",
        # "Firm power output": f"{firm_power_output:,.2f} kW_firm",
        "renewable_share": f"{renewable_share:.2f}%",
        # "Investment per kW_firm": f"{investment_per_kwfirm:,.2f} {currency_symbol}/kW_firm",
        "tariff": f"{tariff:.2f} {currency_symbol}/kWh",
        "project_lifetime": f"{project_lifetime}",
        "bm_name": f"{bm_name}",
        "exchange_rate": f"{project.economic_data.exchange_rate} {currency_symbol}/USD",
        "total_investments": f"{total_investments:,.0f} {currency_symbol}",
    }
    return project_summary


def help_icon(help_text=""):
    return "<a data-bs-toggle='tooltip' title='' data-bs-original-title='{}' data-bs-placement='right'><img style='height: 1.2rem;margin-left:.5rem' alt='info icon' src='{}'></a>".format(
        help_text, static("assets/icons/i_info.svg")
    )


def get_shs_threshold(shs_tier):
    tiers = [tier[0] for tier in HOUSEHOLD_TIERS]
    tiers_verbose = [tier[1] for tier in HOUSEHOLD_TIERS]
    threshold_index = tiers.index(shs_tier)
    excluded_tiers = tiers_verbose[: threshold_index + 1]

    return excluded_tiers


def get_aggregated_cgs(project, as_ts=False):
    options = get_object_or_404(Options, project=project)
    shs_threshold = options.shs_threshold

    # list according to ConsumerType object ids in database
    consumer_types = ["households", "enterprises", "public", "machinery"]
    results_dict = {}

    if len(shs_threshold) != 0:
        shs_consumers = get_shs_threshold(options.shs_threshold)
    else:
        shs_consumers = None

    results_dict["shs"] = {}
    total_demand_shs = 0
    total_consumers_shs = 0

    for consumer_type_id, consumer_type in enumerate(consumer_types, 1):
        results_dict[consumer_type] = {}
        total_demand = np.zeros(8760)
        total_consumers = 0

        # filter consumer group objects based on consumer type
        group_qs = ConsumerGroup.objects.filter(project=project, consumer_type_id=consumer_type_id)

        # calculate total consumers and total demand as sum of array elements in kWh
        for group in group_qs:
            ts = DemandTimeseries.objects.get(pk=group.timeseries_id)

            if shs_consumers is not None and ts.name in shs_consumers:
                total_demand_shs += np.array(ts.get_values_with_unit("kWh")) * group.number_consumers
                total_demand_shs += np.array(ts.get_values_with_unit("kWh")) * group.number_consumers
                total_consumers_shs += group.number_consumers

            else:
                total_demand += np.array(ts.get_values_with_unit("kWh")) * group.number_consumers
                total_consumers += group.number_consumers

        # add machinery total demand to enterprise demand without increasing nr. of consumers
        if consumer_type == "machinery":
            del results_dict[consumer_type]
            results_dict["enterprises"]["total_demand"] += total_demand
        else:
            results_dict[consumer_type]["nr_consumers"] = total_consumers
            results_dict[consumer_type]["total_demand"] = total_demand
            results_dict[consumer_type]["supply_source"] = "mini_grid"

        # TODO SHS demand is manually reset at 0 now to avoid CAPEX and OPEX costs - find better solution
        results_dict["shs"]["nr_consumers"] = total_consumers_shs
        results_dict["shs"]["total_demand"] = np.zeros(8760)
        results_dict["shs"]["supply_source"] = "shs"

    if as_ts is not True:
        for key in results_dict:
            results_dict[key]["total_demand"] = round(sum(results_dict[key]["total_demand"]), 0)

    return results_dict


def get_aggregated_demand(project):
    options = get_object_or_404(Options, project=project)
    shs_threshold = options.shs_threshold
    # Prevent error if no timeseries are present
    total_demand = [np.zeros(8760)]
    cg_qs = ConsumerGroup.objects.filter(project=project)

    if cg_qs.exists():
        # exclude SHS users from aggregated demand for system optimization
        if len(shs_threshold) != 0:
            shs_consumers = get_shs_threshold(options.shs_threshold)
            # TODO need to warn the user if the total_demand is empty due to shs threshold
            cg_qs = cg_qs.exclude(timeseries__name__in=shs_consumers)

        for cg in cg_qs:
            timeseries_values = np.array(cg.timeseries.get_values_with_unit("kWh"))
            nr_consumers = cg.number_consumers
            total_demand.append(timeseries_values * nr_consumers)
        return np.vstack(total_demand).sum(axis=0).tolist()
    else:
        return []


def get_demand_indicators(project, with_timeseries=False):
    """Provide the aggregated, peak and daily averaged demand

    :param with_timeseries: when True return the full timeseries as well
    """
    qs_demand = Asset.objects.filter(scenario=project.scenario, asset_type__asset_type="demand")
    if qs_demand.exists():
        demand = json.loads(qs_demand.get().input_timeseries)
        demand_np = np.array(demand)

    else:
        demand_np = np.array(get_aggregated_demand(project))

    total_demand = demand_np.sum()
    peak_demand = round(demand_np.max(), 1)
    daily_demand = round(total_demand / 365, 1)
    if with_timeseries is True:
        return (demand, total_demand, peak_demand, daily_demand)
    else:
        return (total_demand, peak_demand, daily_demand)


class ReportHandler:
    def __init__(self, project, cache_dict):
        self.doc = Document()
        self.logo_path = "static/assets/logos/cpnigeria-logo.png"
        self.table_counter = 0
        self.figure_counter = 0
        self.heading_counter = 0
        self.subheading_counter = 0

        # set style of the document
        try:
            # Set font "Lato" for the entire self.doc
            self.doc.styles["Normal"].font.name = "Lato"
        except ValueError:
            # Handle the exception when "Lato" is not available and set a fallback font
            self.doc.styles["Normal"].font.name = "Arial"

        # get parameters needed for implementation plan
        options = Options.objects.get(project=project)
        if options.community is not None:
            community_name = options.community.name
        else:
            community_name = project.name
        self.bm = BusinessModel.objects.get(scenario=project.scenario)
        self.bm_name = self.bm.model_name

        # get optimized capacities
        qs_res = FancyResults.objects.filter(simulation__scenario=project.scenario)
        opt_caps = qs_res.filter(
            optimized_capacity__gt=0, asset__in=["pv_plant", "battery", "inverter", "diesel_generator"], direction="in"
        )

        opt_caps = opt_caps.values("asset", "optimized_capacity", "total_flow")

        # Make a sentence with system assets and their optimized capacities
        system_assets = []
        for asset in opt_caps.values_list("asset", "optimized_capacity"):
            if asset[0] != "inverter":
                unit = "kW"
                if asset[0] == "battery":
                    unit = "kWh"
                system_assets.append(f"{asset[0].replace('_', ' ').title()} [{round(asset[1], 1)} {unit}]")
        if len(system_assets) == 1:
            system_assets = "a {}".format(system_assets[0])
        else:
            system_assets = (
                "a {}, " * (len(system_assets) - 2) + "a {}" * (len(system_assets) > 1) + " and a {}"
            ).format(*system_assets)

        self.project = project
        self.cache_dict = cache_dict
        self.image_path = dict(
            es_schema="static/assets/gui/" + options.schema_name,
            bm_graph="static/assets/cp_nigeria/business_models/" + B_MODELS[self.bm_name]["Graph"],
            bm_resp="static/assets/cp_nigeria/business_models/" + B_MODELS[self.bm_name]["Responsibilities"],
        )

        scenario_results = json.loads(KPIScalarResults.objects.get(simulation__scenario=project.scenario).scalar_values)

        lcoe = round(scenario_results["levelized_costs_of_electricity_equivalent"], 2)

        ft = FinancialTool(project)

        total_demand, peak_demand, daily_demand = get_demand_indicators(project)

        inverter_aggregated_flow = ft.system_params.loc[
            (ft.system_params["category"] == "total_flow") & (ft.system_params["supply_source"] == "inverter"), "value"
        ].iloc[0]

        self.aggregated_cgs = get_aggregated_cgs(self.project)

        # Make a sentence with existing enterprises and public facilities
        self.cgs = ConsumerGroup.objects.filter(project=project)
        cg_sentences = {}
        for consumer_type in ["Enterprise", "Public facility"]:
            consumers = []
            # Get the 4 most commonly found consumers for each consumer type and construct into sentence list
            most_consumers = (
                self.cgs.filter(consumer_type__consumer_type=consumer_type)
                .order_by("-number_consumers")
                .values_list("timeseries__name", "number_consumers")[:4]
            )
            for name, nr_consumers in most_consumers:
                consumers.append(f"{name.replace('_', ': ')} ({nr_consumers})")

            cg_sentences[consumer_type] = ", ".join(consumers[:-1]) + f" and {consumers[-1]}"

        self.text_parameters = dict(
            grid_option=options.main_grid,
            bm_name=B_MODELS[self.bm_name]["Verbose"],
            community_name=community_name,
            community_region="Region",
            hh_number=self.aggregated_cgs["households"]["nr_consumers"],
            ent_number=self.aggregated_cgs["enterprises"]["nr_consumers"],
            pf_number=self.aggregated_cgs["public"]["nr_consumers"],
            shs_number=self.aggregated_cgs["shs"]["nr_consumers"],
            shs_threshold=dict(HOUSEHOLD_TIERS)[options.shs_threshold],
            system_assets=system_assets,
            system_capacity="XXXX (???what is the definition of system size???)",
            system_capex=ft.total_capex("NGN"),
            system_opex=ft.total_opex(),
            yearly_production=ft.yearly_production_electricity,
            total_demand=total_demand,
            avg_daily_demand=daily_demand,
            peak_demand=peak_demand,
            diesel_price_increase=ft.financial_params["fuel_price_increase"],
            renewable_share=(inverter_aggregated_flow / total_demand) * 100,
            lcoe=lcoe,
            opex_total=ft.total_opex(),
            opex_growth_rate=ft.opex_growth_rate * 100,
            fuel_costs=ft.fuel_costs,
            fuel_consumption_liter=ft.fuel_consumption_liter,
            energy_system_components_string=options.component_list,
            community_latitude=self.project.latitude,
            community_longitude=self.project.longitude,
            project_name=self.project.name,
            project_lifetime=self.project.economic_data.duration,
            total_investments=ft.total_capex("NGN"),
            disco="DiscoName",
            ent_demand_categories=cg_sentences["Enterprise"],
            pf_demand_categories=cg_sentences["Public facility"],
        )
        if self.text_parameters["grid_option"] is False:
            self.text_parameters["grid_option"] = "is not connected to the national grid"
        else:
            self.text_parameters[
                "grid_option"
            ] = "connected to the national grid, however, the electricity level provided is insufficient"

    def add_heading(self, text, level=1):
        text = text.format(**self.text_parameters).upper()
        if level == 1:
            self.heading_counter += 1
            self.subheading_counter = 0
            number = f"{str(self.heading_counter)}. "
        elif level == 2 and self.heading_counter > 0:
            self.subheading_counter += 1
            number = f"{self.heading_counter}.{self.subheading_counter}. "
        else:
            number = ""

        h = self.doc.add_heading(number + text, level)
        h.runs[0].font.color.rgb = RGBColor(0, 135, 83)
        h.runs[0].font.bold = False
        h.paragraph_format.space_after = Pt(8)

    def add_table(self, rows, cols, caption=None):
        self.table_counter += 1
        t = self.doc.add_table(rows, cols, style="Light Shading")
        if caption is not None:
            self.add_caption(t, caption)

        self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return t

    def add_paragraph(self, text=None, style=None, emph=[]):
        if text is not None:
            try:
                text = text.format(**self.text_parameters)
            except ValueError as e:
                print(text)
                raise (e)
        p = self.doc.add_paragraph(style=style)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        runner = p.add_run(text)
        if "italic" in emph:
            runner.italic = True
        if "bold" in emph:
            runner.bold = True
        if "red" in emph:
            runner.font.color.rgb = RGBColor(255, 0, 0)

        p.paragraph_format.line_spacing = 1.25
        return p

    def add_image(self, path, width=Inches(6), caption=None):
        self.figure_counter += 1
        fig = self.doc.add_picture(path, width=width)
        self.doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if caption is not None:
            self.add_caption(fig, caption)

    def add_image_from_cache(self, name, width=Inches(6), caption=None):
        # the output images available are "cpn_stacked_timeseriesElectricity", "capex", "system_costs", "cash_flow"
        image_data = self.cache_dict.get(name)
        if image_data:
            try:
                image_data = image_data.split(",")[1]
                image_bytes = base64.b64decode(image_data)
                image = io.BytesIO(image_bytes)

                self.add_image(image, width, caption)
            except Exception as e:
                print(e)

    def add_caption(self, tab_or_figure, caption):
        target = {Table: "Table", InlineShape: "Figure"}[type(tab_or_figure)]
        if target == "Table":
            number = self.table_counter
        else:
            number = self.figure_counter

        # caption type
        paragraph = self.doc.add_paragraph(style="Caption")
        paragraph.paragraph_format.space_before = Pt(3)

        # TODO numbering field: code works but only when manually updating the document, replaced with manual numbering for now (see: https://github.com/python-openxml/python-docx/issues/359)
        # run = paragraph.add_run()
        #
        # fldChar = OxmlElement("w:fldChar")
        # fldChar.set(ns.qn("w:fldCharType"), "begin")
        # run._r.append(fldChar)
        #
        # instrText = OxmlElement("w:instrText")
        # instrText.text = f" SEQ {target} \\* ARABIC"
        # run._r.append(instrText)
        #
        # fldChar = OxmlElement("w:fldChar")
        # fldChar.set(ns.qn("w:fldCharType"), "end")
        # run._r.append(fldChar)

        # caption text
        caption_run = paragraph.add_run(f"{target} {number}: {caption}")
        caption_run.font.color.rgb = RGBColor(0, 135, 83)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_dict_as_table(self, dict, caption=None):
        t = self.doc.add_table(len(dict), 2, caption)

        for i, (key, value) in enumerate(dict.items()):
            t.cell(i, 0).text = key
            t.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            t.cell(i, 1).text = value

    def add_df_as_table(self, df, caption=None, index=True):
        if index is True:
            start_idx = 1
        else:
            start_idx = 0

        t = self.add_table(df.shape[0] + 1, df.shape[1] + start_idx, caption)

        # add indices
        if index is True:
            for j in range(df.shape[0]):
                t.cell(j + 1, 0).text = df.index[j]

        # add headers
        for j in range(df.shape[1]):
            t.cell(0, j + start_idx).text = df.columns[j]

        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j + start_idx).text = str(df.values[i, j])

    def get_df_from_cache(self, name):
        # the output tables available are "cost_table", "capex_table", "summary_table", "system_table"
        table_data = self.cache_dict.get(name)
        table_df = pd.DataFrame.from_dict(table_data["data"], orient="index", columns=table_data["headers"])
        return table_df

    def add_table_from_records(self, records, columns=None):
        tot_rows = len(records)
        col_spacer = 0
        if columns is not None:
            tot_rows += 1
            col_spacer = 1

        table = self.add_table(rows=tot_rows, cols=len(records[0]))

        if columns is not None:
            hdr_cells = table.rows[0].cells
            for i, item in enumerate(columns):
                hdr_cells[i].text = item.format(**self.text_parameters)

        for j, record in enumerate(records):
            row_cells = table.rows[j + col_spacer].cells
            # row_cells = table.add_row().cells
            for i, item in enumerate(record):
                if isinstance(item, str):
                    try:
                        row_cells[i].text = item.format(**self.text_parameters)
                    except ValueError as e:
                        print(item)
                        raise (e)
                else:
                    row_cells[i].text = str(item)

    def add_financial_table(self, records, title=""):
        table = self.add_table(rows=1, cols=2)
        row_cells = table.rows[0].cells
        row_cells[0].merge(row_cells[-1])

        # Set a cell background (shading) color and align center
        shading_elm = parse_xml(r'<w:shd {} w:fill="008753"/>'.format(ns.nsdecls("w")))
        row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        row_cells[0].text = title
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for j, record in enumerate(records):
            row_cells = table.add_row().cells
            for i, item in enumerate(record):
                try:
                    row_cells[i].text = item.format(**self.text_parameters)
                except ValueError as e:
                    print(item)
                    raise (e)

        # TODO add a thick line

        row_cells = table.add_row().cells
        row_cells[0].text = "Total"
        row_cells[0].bold = True

    def add_list(self, list_items, style=None, emph=[]):
        if isinstance(list_items, str):
            list_items = [list_items]
        if style is None:
            style = "List Bullet"
        for bullet in list_items:
            self.add_paragraph(bullet, style=style, emph=emph)

    def create_attribute(self, element, name, value):
        # used for page numbering
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        # code from https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx
        fldChar1 = OxmlElement("w:fldChar")
        self.create_attribute(fldChar1, "w:fldCharType", "begin")

        instrText = OxmlElement("w:instrText")
        self.create_attribute(instrText, "xml:space", "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        self.create_attribute(fldChar2, "w:fldCharType", "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def prevent_table_splitting(self):
        tags = self.doc.element.xpath("//w:tr[position() < last()]/w:tc/w:p")
        for tag in tags:
            ppr = tag.get_or_add_pPr()
            ppr.keepNext_val = True

    def add_hyperlink_into_run(self, paragraph, run, url):
        # code from: https://github.com/python-openxml/python-docx/issues/74#issuecomment-441351994
        runs = paragraph.runs
        for i in range(len(runs)):
            if runs[i].text == run.text:
                break

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement("w:hyperlink")
        self.create_attribute(hyperlink, "r:id", r_id)
        hyperlink.append(run._r)
        paragraph._p.insert(i + 1, hyperlink)

    def add_paragraph_with_hyperlink(self, paragraph_text, hyperlink_text, url):
        p = self.add_paragraph()

        # Add the first run (text before the hyperlink)
        p.add_run(paragraph_text.split(hyperlink_text)[0])

        # Add the second run (hyperlink text)
        run_hyperlink = p.add_run(hyperlink_text)
        run_hyperlink.bold = True
        run_hyperlink.underline = True
        run_hyperlink.font.color.rgb = RGBColor(0, 102, 204)
        self.add_hyperlink_into_run(p, run_hyperlink, url)

        # Add the third run (text after the hyperlink)
        p.add_run(paragraph_text.split(hyperlink_text)[1])

        return p

    @staticmethod
    def create_community_criteria_list(bmanswer_qs):
        BM_CRITERIA = csv_to_dict("business_model_report_criteria.csv", label_col="Criteria")
        criteria_list = []
        for criteria, values in BM_CRITERIA.items():
            total_score_qs = bmanswer_qs.filter(question_id__in=json.loads(values["Questions"])).aggregate(
                total_score=Sum("score")
            )
            score = total_score_qs["total_score"]
            if score > json.loads(values["Threshold"]):
                criteria_list.append(criteria)

        return criteria_list

    def save(self, response):
        self.doc.save(response)

    def create_cover_sheet(self):
        title = f"Mini-Grid Implementation Plan for {self.project.name}"
        subtitle = f"Date: {date.today().strftime('%d.%m.%Y')}"
        summary = f"{self.project.description}"

        # Add logo in the right top corner
        header = self.doc.sections[0].header
        run = header.paragraphs[0].add_run()
        run.add_picture(self.logo_path, width=Pt(70))
        header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        header.paragraphs[0].paragraph_format.space_after = Pt(3)

        # Add title
        title_paragraph = self.add_paragraph()
        title_paragraph.paragraph_format.space_before = Inches(0.5)
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add subtitle
        subtitle_paragraph = self.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(13)
        subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add project details
        self.add_heading("Project details", level=2)

        # Add project summary table
        project_summary = get_project_summary(self.project)
        summary_table = {}
        for key in project_summary:
            summary_table[OUTPUT_PARAMS[key]["verbose"]] = project_summary[key]

        self.add_dict_as_table(summary_table)

        # Add project summary
        self.add_heading("Project summary", level=2)
        self.add_paragraph(
            "The {community_name} community is a rural community in the {community_region} region. The community "
            "comprises about {hh_number} households as well as {ent_number} enterprises and {pf_number} public "
            "facilities. Currently, the community {grid_option}. In light of the community's aspiration to achieve a "
            "constant and reliable electricity supply, the community is willing to engage with project partners to "
            "construct a suitable mini-grid system. The system proposed in this implementation plan, by using the "
            "CP-Nigeria Toolbox, comprises {system_assets}. Overall Capex would "
            "amount to {system_capex:,.0f} NGN, while overall Opex amount to {system_opex:,.0f} NGN. Regarding the "
            "project implementation, an {bm_name} business model approach is suggested."
        )

        subtitle_run = subtitle_paragraph.add_run(summary)
        subtitle_run.font.size = Pt(10)
        subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Add page break
        self.doc.add_page_break()

    def create_report_content(self):
        # TODO implement properly

        ### INTRODUCTION ###
        self.add_heading("Context and Background")
        self.add_heading("Community Background", level=2)
        self.add_list(
            (
                "The community is located in the {community_region} region",
                "The community comprises about {hh_number} households as well as {ent_number} enterprises and {pf_number} public facilities.",
                "Some of the most commonly found enterprises in the community include {ent_demand_categories}",
                "Existing public facilities include {pf_demand_categories}",
                "Currently, the community is {grid_option}",
            )
        )

        # self.add_table(
        #     (
        #         ("Consumption Level", "Very Low", "Low", "Middle", "High", "Very High"),
        #         ("Number of households", "XX", "XX", "XX", "XX", "XX"),
        #     )
        # )

        self.add_paragraph(
            "--------------------------------------\nDear toolbox user, it is recommended to enrich this section by providing the following information",
            emph=["italic", "red"],
        )
        self.add_list("Description of the current electricity access situation", emph=["italic", "red"])
        self.add_list(
            ("Use of diesel generators", "Most common cooking practices", "Use of lighting"),
            style="List Bullet 2",
            emph=["italic", "red"],
        )
        self.add_list(
            (
                "Experiences with community-led projects",
                "Existing community organizational structures, e.g. (energy) cooperatives, associations, committees, etc.",
                "Skills of community members in the context of energy/electricity supply",
            ),
            emph=["italic", "red"],
        )

        self.add_heading("CP-Nigeria Toolbox Context", level=2)
        self.add_paragraph_with_hyperlink(
            "This implementation plan is the output of an open source online toolbox.",
            "open source online toolbox.",
            "https://community-minigrid.ng/en/",
        )
        self.add_paragraph_with_hyperlink(
            "The toolbox has been "
            "developed within the frame of the project CP Nigeria: Communities of Practice as driver of a bottom-up "
            "energy transition in Nigeria: This project is funded by the International Climate Initiative (IKI). "
            "The overall objective of the project is to achieve a climate-friendly energy supply through "
            "decentralized renewable energy (DRE) in Nigeria by 2030. In order to achieve this, civil society must "
            "play a driving role. The project therefore aims to create exemplary civil society nuclei ('Communities "
            "of Practice') and to empower them to plan and implement DRE projects (local level). Based on this, "
            "it works transdisciplinary on improving the political framework for local decentralized RE projects ("
            "national level). In this way, civil society can be empowered to implement DRE independently and make a "
            "significant contribution to the achievement of climate change mitigation goals.  ",
            "International Climate Initiative (IKI)",
            "https://www.international-climate-initiative.com/en/project/communities-of-practice-as-driver-of-a-bottom-up-energy-transition-in-nigeria-img2020-i-003-nga-energy-transition-communities-of-practice/",
        )

        ### DEMAND SECTION ###
        self.add_heading("Electricity Demand Profile")

        p = self.add_paragraph_with_hyperlink(
            "The total electricity demand is estimated by assigning demand profiles to the households, enterprises "
            "and public facilities present in the community, based on demand profiles constructed within the "
            "PeopleSun project, which were derived from surveys and appliance audits conducted within "
            "electrified communities. As the demand is based on proxy data for already electrified communities, "
            "the system doesn't consider any demand increase over project lifetime, but instead assumes that the "
            "proposed supply system would be able to satisfy future demand.",
            "PeopleSun project",
            "https://www.peoplesun.org/",
        )

        self.add_paragraph(
            "Given that not all households may be connected to the mini-grid, but some may be served by Solar Home "
            "Systems (SHS) instead, all households assigned to the {shs_threshold} tier and below are excluded from the "
            "supply system optimization. In this case, {shs_number} households were assumed to be served by SHS."
        )

        self.add_paragraph(
            "In total, the community has {hh_number} households, {ent_number} enterprises and {pf_number} public "
            "facilities that would be connected to the mini-grid. Due to the nature of the used demand profiles, enterprise profiles only include basic "
            "amenities (e.g. lighting), while heavy machinery is added separately. The following table displays the "
            "number of consumers and their respective yearly electricity demand."
        )

        self.add_df_as_table(self.get_df_from_cache("summary_table"), caption="Community demand summary")

        self.add_paragraph(
            "The total estimated yearly demand for the {community_name} community is {total_demand:,.0f} kWh/year. The "
            "average daily demand is {avg_daily_demand:,.1f} kWh/day, while the peak demand for the simulated year is "
            "{peak_demand:,.1f} kW. The following graph displays how the cumulated demand is aggregated based on the given community "
            "characteristics."
        )

        self.add_image_from_cache("mini_grid_demand", caption="Total mini-grid demand")

        ### ELECTRICITY SYSTEM SECTION ###
        self.add_heading("Electricity Supply System Size and Composition")
        self.add_image(self.image_path["es_schema"], width=Inches(3), caption="System schematic")

        self.add_paragraph(
            "Based on the calculated yearly demand, a least-cost-optimization was conducted for a supply system with "
            "the following components: {energy_system_components_string}. The cost and asset characteristics used to "
            "conduct the system optimization can be seen in the annex. Based on the given system setup, the following "
            "asset sizes result in the least-cost solution:"
        )

        # TODO put this in annex
        # - {Table with supply system parameters (form contents from page 4)}

        # Table with optimized capacities
        self.add_df_as_table(self.get_df_from_cache("system_table"), caption="System size")

        self.add_paragraph(
            "The system presents a levelized cost of electricity (LCOE) of {lcoe:.2f} NGN/kWh, "
            "with {renewable_share:.2}% of the generation coming from renewable sources."
        )

        self.add_paragraph(
            "In total, the investment costs for the power supply system amount to {system_capex:,.0f} NGN. More detailed "
            "information regarding investment costs and financial considerations can be found in the respective section. The "
            "operational expenditures for the simulated year amount to {opex_total:,.0f} NGN, with an estimated annual "
            "increase in operational expenditures of {opex_growth_rate}%. Of these expenditures, {fuel_costs:,.0f} NGN are "
            "attributed to fuel costs, amounting to {fuel_consumption_liter:,.1f} L consumed. The "
            "following graph displays the power flow for the system during one week:"
        )

        # Stacked timeseries graph
        self.add_image_from_cache("cpn_stacked_timeseriesElectricity", caption="Power flows during first week")

        ### BUSINESS MODEL SECTION ###
        self.add_heading("Business Model of the Mini-grid Project")

        self.add_paragraph(B_MODELS[self.bm_name]["Report"])

        self.add_image(self.image_path["bm_graph"], width=Inches(2.5), caption="Business model structure")
        self.add_image(self.image_path["bm_resp"], width=Inches(5), caption="Business model responsibilities")

        # Add additional text and criteria bullet points if community-led model is chosen
        if "cooperative" in self.bm_name:
            self.add_paragraph(
                "A cooperative-led model is an innovative approach in Nigeria that strongly enhances "
                "local awareness and engagement. Additionally, the community-led approach can increase "
                "the affordability of tariffs for customers in the community. It is, however, "
                "challenging for communities to attract adequate financial resources, therefore, the "
                "community is reaching out to financial partners."
            )

            qs = BMAnswer.objects.filter(business_model=self.bm)
            if qs.exists():
                self.add_paragraph(
                    "Within the process of the CP-Nigeria Toolbox, the community’s suitability for following "
                    "a community-led approach for the realisation of a mini-grid system has been assessed "
                    "based on several research-based criteria. The community meets the following criteria "
                    "to choose the cooperative-led model and has in place:"
                )
                self.add_list(self.create_community_criteria_list(qs))

        ### FINANCIAL ANALYSIS SECTION ###
        self.add_heading("Financial Analysis")
        self.add_heading("Capital Expenditure (CAPEX)", level=2)

        self.add_paragraph(
            "The project includes the following Capital Expenditures (CAPEX) that need to be covered "
            "during installation of the technical equipment, i.e. the power supply system and the "
            "transmission network. CAPEX are considered in a conservative way and include also "
            "logistical costs, an insurance for construction as well as planning and labor costs. VAT is"
            " considered for non-technical equipment costs only."
        )

        self.add_df_as_table(self.get_df_from_cache("capex_table"), caption="Total mini-grid CAPEX")
        self.add_image_from_cache("capex", caption="Total mini-grid CAPEX")
        # self.add_financial_table((("", ""), ("", "")), title="Capex (in USD)")

        self.add_heading("Operational Expenditure (OPEX)", level=2)
        self.add_paragraph(
            "The project includes the following operational expenditures (OPEX) that need to be covered"
            " during operation of the power supply system and the transmission network. OPEX are "
            "considered in a conservative way and include service and maintenance for the mini-grid as "
            "well as diesel fuel for the diesel-genset."
        )
        self.add_df_as_table(self.get_df_from_cache("cost_table"), "Total system costs during first year")
        # self.add_financial_table((("", ""), ("", "")), title="Opex (in USD)")

        self.add_heading("Key Financial Parameters", level=2)
        self.add_paragraph(
            "The following key financial parameters describe the economic viability and profitability "
            "of the project. They include the resulting community tariff and its annual increase rate "
            "as well as the internal rate of return (IRR) of the project activity over 10 and 20 years."
        )

        self.add_df_as_table(self.get_df_from_cache("financial_kpi_table"), caption="Key Financial Parameters")

        self.add_heading("Financing Structure", level=2)
        self.add_paragraph(
            "The financing structure shows the key financial conditions including the communities’ and "
            "mini-grid companies’ equity share and interest, the grant and the debt volume and it’s "
            "average interest rate. For the grant, a deduction of 25% is considered, since typically "
            "performance based grants are applied that are provided after the implementation of the "
            "project. Thus a domestic bank loan with approximately 25% interest has to be taken up for "
            "one year, effectively reducing the volume of the grant by 25%. WACC describes the "
            "resulting weighted costs of capital for the project activity."
        )

        self.add_df_as_table(
            self.get_df_from_cache("financing_structure_table"), caption="Financing structure for the project"
        )

        self.add_heading("Cash Flow Diagram", level=2)
        self.add_paragraph(
            "The following Figure shows the net cash flow over time, including debt repayment and debt "
            "interest payments as well as the  comparison of net operating revenues and  operating expenses."
        )
        self.add_image_from_cache("cash_flow", caption="Cash flow over project lifetime")
        # TODO add next steps
        self.add_heading("Next Steps")
        self.doc.add_page_break()
        self.add_heading("Appendix")
        # TODO add tables about cost assumptions etc here


class FinancialTool:
    cost_assumptions = pd.read_csv(staticfiles_storage.path("financial_tool/cost_assumptions.csv"), sep=";")
    loan_assumptions = {"Tenor": 10, "Grace period": 1, "Cum. replacement years": 10}

    def __init__(self, project):
        """
        The financial tool object should be initialized with the opt_caps and asset_costs objects created in the results
        page. When a financial tool object is initialized, the project start and duration are set, the relevant economic
        inputs and simulation results are fetched and the system size over the project lifetime is calculated
        (considering potential demand and consumer increase (both mini-grid and SHS over project lifetime)
        """
        # TODO there are a number of loose variables (unclear if default or missing in tool) - see list in PR and
        #  discuss along with best approach to display results
        self.project = project
        self.exchange_rate = project.economic_data.exchange_rate
        self.project_start = project.scenario.start_date.year
        self.project_duration = project.economic_data.duration

        self.financial_params = self.collect_financial_params()
        self.system_params = self.collect_system_params()
        # calculate the system growth over the project lifetime and add rows for new mg and shs consumers per year
        system_lifetime = self.growth_over_lifetime_table(
            self.system_params[self.system_params["category"].isin(["nr_consumers", "total_demand"])],
            "value",
            "growth_rate",
            "label",
        )
        self.system_lifetime = self.add_diff_rows(system_lifetime)

    def collect_system_params(self):
        """
        This method takes the optimized capacities and cost results as inputs and returns a dataframe with all the
        relevant system and cost parameters for the results (capex, opex, replacement costs and asset sizes). The
        fetched diesel OM costs are reset to the original diesel price and the diesel price increase is set as the
        growth rate (this is not optimal since it is essentially "faking" the simulation results). A label is set for
        each system param that corresponds to the target set in cost_assumptions, e.g. the target for PV installation
        labor will be the PV optimized capacity, which is set as a label in this dataframe to later merge in
        calculate_capex().
        """
        growth_rate = 0.0

        # get capacities and cost results
        qs_res = FancyResults.objects.filter(simulation__scenario=self.project.scenario)
        opt_caps = qs_res.filter(
            optimized_capacity__gt=0, asset__in=["pv_plant", "battery", "inverter", "diesel_generator"], direction="in"
        ).values("asset", "optimized_capacity", "total_flow")
        qs_installed = Asset.objects.filter(
            scenario=self.project.scenario,
            asset_type__asset_type__in=["pv_plant", "capacity", "diesel_generator", "transformer_station_in"],
        ).values("asset_type__asset_type", "installed_capacity")
        # inst_caps = {'battery' if item['asset_type__asset_type'] == 'capacity'
        #              else item['asset_type__asset_type']: item['installed_capacity'] for item in qs_installed}
        inst_caps = {}
        for asset in qs_installed:
            if asset["asset_type__asset_type"] == "capacity":
                asset_name = "battery"
            elif asset["asset_type__asset_type"] == "transformer_station_in":
                asset_name = "inverter"
            else:
                asset_name = asset["asset_type__asset_type"]
            inst_caps[asset_name] = asset["installed_capacity"]

        costs = get_costs(self.project.scenario.simulation)
        costs["opex_total"] = costs["opex_var_total"] + costs["opex_fix_total"]
        costs.drop(columns=["opex_var_total", "opex_fix_total", "capex_total"], inplace=True)

        total_demand = (
            pd.DataFrame.from_dict(get_aggregated_cgs(self.project), orient="index").groupby("supply_source").sum()
        )
        asset_sizes = pd.DataFrame.from_records(opt_caps).set_index("asset")
        asset_sizes["installed_capacity"] = inst_caps
        assets = pd.concat([asset_sizes, costs], axis=1)

        system_params = pd.concat(
            [
                pd.melt(assets.reset_index(), id_vars=["index"], var_name=["category"]).rename(
                    columns={"index": "supply_source"}
                ),
                pd.melt(total_demand.reset_index(), id_vars=["supply_source"], var_name=["category"]),
            ],
            ignore_index=True,
        )
        # TODO replace this with custom consumer/demand increase when it is included in tool GUI - maybe not prio since
        #  we are already considering a future demand scenario
        system_params["growth_rate"] = growth_rate
        system_params.loc[system_params["category"] == "opex_total", "growth_rate"] = self.opex_growth_rate
        system_params.loc[system_params["category"] == "fuel_costs_total", "growth_rate"] = self.financial_params[
            "fuel_price_increase"
        ]
        system_params["label"] = system_params["supply_source"] + "_" + system_params["category"]
        # TODO include excess generation (to be used by excess gen tariff - also not prio while not considering feedin)
        return system_params

    @property
    def yearly_production_electricity(self):
        flow_df = self.system_params[self.system_params["category"] == "total_flow"]
        return flow_df[flow_df["supply_source"].isin(["diesel_generator", "pv_plant", "dso"])].value.sum()

    @property
    def fuel_costs(self):
        return self.system_params[
            (self.system_params["supply_source"] == "diesel_generator")
            & (self.system_params["category"] == "fuel_costs_total")
        ].value.iloc[0]

    @property
    def fuel_consumption_liter(self):
        return (
            self.system_params[
                (self.system_params["supply_source"] == "diesel_generator")
                & (self.system_params["category"] == "total_flow")
            ].value.iloc[0]
            / ENERGY_DENSITY_DIESEL
        )

    def collect_financial_params(self):
        """
        This method collects the financial parameters from the economic and the equity data set on the project.
        """
        qs_eq = EquityData.objects.filter(scenario=self.project.scenario).values(
            "debt_start",
            "fuel_price_increase",
            "grant_share",
            "debt_interest_MG",
            "debt_interest_replacement",
            "debt_interest_SHS",
            "loan_maturity",
            "grace_period",
            "equity_interest_MG",
            "equity_interest_SHS",
            "equity_community_amount",
            "equity_developer_amount",
        )
        qs_ed = EconomicData.objects.filter(project=self.project).values("discount", "tax")
        financial_params = qs_ed.first() | qs_eq.first()
        financial_params["capex_fix"] = self.project.scenario.capex_fix

        return financial_params

    @staticmethod
    def yearly_increase(base_amount, growth_rate, year):
        return base_amount * (1 + growth_rate) ** year

    def growth_over_lifetime_table(self, df, base_col, growth_col, index_col=None):
        """
        This method creates a wide table over the project lifetime based on start values and growth rate of the given
        parameters. It is used to calculate both the system growth over the project lifetime and the growth in
        costs/revenues based on a given annual increase rate (that can be set independently for each parameter in
        collect_system_params() and cost_assumptions.csv). This method is used in the revenue and losses_over lifetime
        methods.
        """
        lifetime_df = pd.DataFrame(
            {
                self.project_start
                + year: self.yearly_increase(base_amount=df[base_col], growth_rate=df[growth_col], year=year)
                for year in range(self.project_duration)
            }
        )

        if index_col is not None:
            if isinstance(index_col, list):
                # Create a MultiIndex with the provided list
                lifetime_df.index = [np.array(df[col]) for col in index_col]
            else:
                lifetime_df.index = df[index_col]

        return lifetime_df

    def add_diff_rows(self, df):
        """
        This method adds new consumers rows over project lifetime both for mini-grid and SHS consumers. This is needed
        to calculate the connection fees for new consumers per project year, given a yearly consumer increase. So far,
        consumer increase is not accounted for, so all but the initial year should return 0.
        """
        vars = [ix for ix in df.index if "nr_consumers" in ix]
        for var in vars:
            diff = df.T[var].diff()
            diff[self.project_start] = df[self.project_start].loc[var]
            new_var = f"{var}_new"
            df.loc[new_var] = diff

        return df

    @property
    def capex(self):
        """
        This method takes the given general cost assumptions and merges them with the specific project results
        (asset sizes) by target variable. The system costs from the simulation are appended to the dataframe and the
        VAT tax is calculated. The method returns a dataframe with all CAPEX costs by category.
        """
        capex_df = pd.merge(
            self.cost_assumptions[~self.cost_assumptions["Category"].isin(["Revenue", "Solar home systems"])],
            self.system_params[["label", "value"]],
            left_on="Target",
            right_on="label",
            suffixes=("_costs", "_outputs"),
            how="left",
        )

        capex_df["Total costs [USD]"] = capex_df["USD/Unit"] * capex_df["value"]

        vat_costs = (
            capex_df.groupby("Category")["Total costs [USD]"].sum()[["Logistics", "Labour and soft costs"]].sum()
        )

        vat_percentage = self.financial_params["tax"]
        capex_df.loc[len(capex_df)] = {
            "Description": "VAT",
            "Category": "Taxes",
            "Qty": 1,
            "Total costs [USD]": vat_costs * vat_percentage,
        }
        capex_df["Total costs [NGN]"] = capex_df["Total costs [USD]"] * self.exchange_rate

        # add fix project costs from economic data view to capex dataframe
        capex_df.loc[len(capex_df)] = {
            "Description": "Planning and development costs",
            "Category": "Fix project costs",
            "Qty": 1,
            "Total costs [NGN]": self.financial_params["capex_fix"],
        }

        # Add the power supply system capex from the sim results (accounts for custom cost parameters instead of static)
        system_capex_rows = [
            {
                "Description": row["supply_source"].replace("_", " ").title(),
                "Category": "Power supply system",
                "Total costs [NGN]": row["value"],
            }
            for index, row in self.system_params[self.system_params["category"] == "capex_initial"].iterrows()
        ]

        capex_df = pd.concat([capex_df, pd.DataFrame(system_capex_rows)], ignore_index=True)

        return capex_df  # capex_df['Total costs [USD]'].sum() returns gross CAPEX

    def total_capex(self, currency="NGN"):
        """Sum of all capex costs"""
        print("capex table\t", self.capex[f"Total costs [{currency}]"].sum())
        capex_df = self.system_params[self.system_params["category"].isin(["capex_initial", "capex_replacement"])]
        print("system params\t", capex_df.value.sum())
        return self.capex[f"Total costs [{currency}]"].sum()

    def total_opex(self):
        costs_om_df = self.system_params[self.system_params["category"] == "opex_total"]
        return costs_om_df.value.sum()

    @property
    def opex_growth_rate(self):
        return self.cost_assumptions.loc[self.cost_assumptions["Category"] == "Opex", "Growth rate"].iloc[0]

    def revenue_over_lifetime(self, custom_tariff=None):
        """
        This method returns a wide table calculating the revenue flows over project lifetime based on the cost
        assumptions for revenue together with the number of consumers and total demand of the system.
        """
        # set the tariff to the already calculated tariff if the function is not being called from goal seek
        if custom_tariff is not None:
            self.cost_assumptions.loc[
                self.cost_assumptions["Description"] == "Community tariff", "USD/Unit"
            ] = custom_tariff

        revenue_df = pd.merge(
            self.cost_assumptions[self.cost_assumptions["Category"] == "Revenue"],
            self.system_params[["label", "value", "growth_rate"]],
            left_on="Target",
            right_on="label",
            suffixes=("_costs", "_outputs"),
            how="left",
        )

        revenue_lifetime = (
            self.growth_over_lifetime_table(
                revenue_df, "USD/Unit", growth_col="Growth rate", index_col=["Description", "Target"]
            )
            * self.exchange_rate
        )

        # here level is set to 1 because above "Target" column of the revenue_df is on level 1 of the MultiIndex of revenue_lifetime
        revenue_flows = revenue_lifetime.mul(self.system_lifetime, level=1)
        revenue_flows.loc[("Total operating revenues", "operating_revenues_total"), :] = revenue_flows.sum()

        return revenue_flows

    @property
    def om_costs_over_lifetime(self):
        """
        This method returns a wide table calculating the OM cost flows over project lifetime based on the OM costs of
        the system together with the annual cost increase assumptions.
        """

        costs_om_df = self.system_params[self.system_params["category"].isin(["opex_total", "fuel_costs_total"])]
        costs_om_df = costs_om_df[costs_om_df["value"] != 0]
        om_lifetime = self.growth_over_lifetime_table(costs_om_df, "value", growth_col="growth_rate", index_col="label")

        # TODO include these costs in extra information about solar home systems but not general mini-grid
        # shs_costs_om = (
        #     self.cost_assumptions.loc[
        #         self.cost_assumptions["Description"] == "SHS operational expenditures", "USD/Unit"
        #     ].values[0]
        #     * self.exchange_rate
        # )

        # multiply the unit prices by the amount
        # om_lifetime.loc["opex_total_shs"] = self.system_lifetime.loc["shs_nr_consumers"] * shs_costs_om

        # calculate total operating expenses
        om_lifetime.loc["opex_total"] = om_lifetime.sum()

        return om_lifetime

    def debt_service_table(self, amount, tenor, gp, ir, debt_start):
        debt_service = pd.DataFrame(
            index=["Interest", "Principal", "Balance opening", "Balance closing", "Capital service"],
            columns=range(self.project_start - 1, self.project_start + self.project_duration),
            data=0.0,
        )
        debt_service[debt_start - 1] = [0, 0, amount, amount, 0]

        for index, year in enumerate(range(debt_start, debt_start + tenor), 1):
            per = index - gp
            nper = tenor - gp
            if index <= tenor:
                if index > gp:
                    debt_service.at["Principal", year] = -npf.ppmt(ir, per, nper, amount)
                else:
                    debt_service.at["Principal", year] = 0
                debt_service.at["Balance opening", year] = debt_service.loc["Balance closing", year - 1]
                debt_service.at["Interest", year] = debt_service.loc["Balance opening", year] * ir
                debt_service.at["Balance closing", year] = (
                    debt_service.loc["Balance opening", year] - debt_service.at["Principal", year]
                )
                debt_service.at["Capital service", year] = (
                    debt_service.loc["Interest", year] + debt_service.at["Principal", year]
                )
            else:
                debt_service.at[:, year] = 0

        return debt_service

    @property
    def initial_loan_table(self):
        """
        This method creates a table for the initial CAPEX debt according to the debt share (CAPEX - grant and equity).
        """
        tenor = self.financial_params["loan_maturity"]
        grace_period = self.financial_params["grace_period"]
        amount = self.financial_kpis["initial_loan_amount"]
        interest_rate = self.financial_params["debt_interest_MG"]
        debt_start = self.project_start

        return self.debt_service_table(
            amount=amount, tenor=tenor, gp=grace_period, ir=interest_rate, debt_start=debt_start
        )

    @property
    def replacement_loan_table(self):
        """
        This method creates a table for the replacement costs debt (accounts for replacing battery, inverter and diesel
        generator after 10 years).
        """
        # TODO should this always be 10 or depending on given shortest component lifetime
        tenor = self.financial_params["loan_maturity"]
        debt_start = self.project_start + self.loan_assumptions["Cum. replacement years"]
        grace_period = self.financial_params["grace_period"]
        amount = self.financial_kpis["replacement_loan_amount"]
        interest_rate = self.financial_params["debt_interest_replacement"]

        return self.debt_service_table(
            amount=amount, tenor=tenor, gp=grace_period, ir=interest_rate, debt_start=debt_start
        )

    def losses_over_lifetime(self, custom_tariff=None):
        """
        This method first calculates the EBITDA (earnings before interest, tax, depreciation and amortization), then
        the financial losses through depreciation, interest payments to get the EBT (earnings before tax) and finally
        including taxes to get the net income over the project lifetime. The calculate_capex() method is used here.
        """
        depreciation_yrs = self.project_duration
        capex_df = self.capex
        system_capex = capex_df[capex_df["Category"] == "Power supply system"]["Total costs [NGN]"].sum()
        equity = self.financial_params["equity_community_amount"] + self.financial_params["equity_developer_amount"]
        losses = pd.DataFrame(columns=range(self.project_start, self.project_start + self.project_duration))

        losses.loc["EBITDA"] = (
            self.revenue_over_lifetime(custom_tariff).loc[("Total operating revenues", "operating_revenues_total"), :]
            - self.om_costs_over_lifetime.loc["opex_total"]
        )
        losses.loc["Depreciation"] = 0.0
        losses.loc["Depreciation"][:depreciation_yrs] = system_capex / depreciation_yrs
        losses.loc["Equity interest"] = equity * self.financial_params["equity_interest_MG"]
        losses.loc["Debt interest"] = (
            self.initial_loan_table.loc["Interest"] + self.replacement_loan_table.loc["Interest"]
        )
        losses.loc["Debt repayments"] = (
            self.initial_loan_table.loc["Principal"] + self.replacement_loan_table.loc["Principal"]
        )

        losses.loc["EBT"] = (
            losses.loc["EBITDA"]
            - losses.loc["Depreciation"]
            - losses.loc["Equity interest"]
            - losses.loc["Debt interest"]
        )
        losses.loc["Corporate tax"] = [
            losses.loc["EBT"][year] * self.financial_params["tax"] if losses.loc["EBT"][year] > 0 else 0
            for year in losses.columns
        ]
        losses.loc["Net income"] = losses.loc["EBT"] - losses.loc["Corporate tax"]

        return losses

    def cash_flow_over_lifetime(self, custom_tariff=None):
        """
        This method calculates the cash flows over system lifetime considering the previously calculated loan debt,
        earnings and financial losses. Both the losses_ver_lifetime() and debt_service_table() are called here.
        """
        cash_flow = pd.DataFrame(columns=range(self.project_start, self.project_start + self.project_duration))
        losses = self.losses_over_lifetime(custom_tariff)

        cash_flow.loc["Cash flow from operating activity"] = losses.loc["EBITDA"] - losses.loc["Corporate tax"]
        cash_flow.loc["Cash flow after debt service"] = (
            cash_flow.loc["Cash flow from operating activity"]
            - losses.loc["Equity interest"]
            - losses.loc["Debt interest"]
            - losses.loc["Debt repayments"]
        )
        cash_flow.loc["Free cash flow available"] = (
            losses.loc["EBITDA"]
            - losses.loc["Corporate tax"]
            - losses.loc["Equity interest"]
            - losses.loc["Debt interest"]
            - losses.loc["Debt repayments"]
        )

        return cash_flow

    def internal_return_on_investment(self, years):
        """
        This method returns the internal return on investment % based on project CAPEX, grant share and cash flows
        after a given number of project years.
        """
        gross_capex = self.capex["Total costs [NGN]"].sum()
        grant = self.financial_params["grant_share"] * gross_capex
        cash_flow = self.cash_flow_over_lifetime()
        cash_flow_irr = cash_flow.loc["Cash flow from operating activity"].tolist()
        cash_flow_irr.insert(0, -gross_capex + grant)

        return npf.irr(cash_flow_irr[: (years + 1)])

    def goal_seek_helper(self, custom_tariff):
        cashflow_helper = np.sum(
            self.cash_flow_over_lifetime(custom_tariff).loc["Cash flow after debt service"].tolist()[0:5]
        )
        return cashflow_helper

    @property
    def financial_kpis(self):
        usable_grant = 0.75  # this factor assumes that 25% of the grant is directly used for loan interest payments
        gross_capex = self.capex["Total costs [NGN]"].sum()
        total_equity = (
            self.financial_params["equity_community_amount"] + self.financial_params["equity_developer_amount"]
        )
        total_grant = self.financial_params["grant_share"] * gross_capex * usable_grant
        initial_amount = max(gross_capex - total_grant - total_equity, 0)
        replacement_amount = self.capex[self.capex["Description"].isin(["Battery", "Inverter", "Diesel Generator"])][
            "Total costs [NGN]"
        ].sum()

        loan_fraction = initial_amount / gross_capex
        equity_fraction = total_equity / gross_capex
        wacc = (loan_fraction * self.financial_params["debt_interest_MG"]) + (
            equity_fraction * self.financial_params["equity_interest_MG"]
        )

        financial_kpis = {
            "total_investments": gross_capex,
            "equity_community": self.financial_params["equity_community_amount"],
            "equity_developer": self.financial_params["equity_developer_amount"],
            "total_equity": total_equity,
            "initial_loan_amount": initial_amount,
            "interest_rate": self.financial_params["debt_interest_MG"],
            "maturity": self.financial_params["loan_maturity"],
            "grace_period": self.financial_params["grace_period"],
            "replacement_loan_amount": replacement_amount,
            "total_grant": total_grant,
            "wacc": wacc,
        }

        return financial_kpis

    @property
    def tariff(self):
        x = np.arange(0.2, 0.5, 0.2)
        # compute the sum of the cashflow for the first 4 years for different tariff (x)
        # as this is a linear function of the tariff, we can fit it and then find the tariff value x0
        # for which the sum of the cashflow for the first 4 years is 0
        m, h = np.polyfit(x, [self.goal_seek_helper(xi) for xi in x], deg=1)
        x0 = -h / m
        return x0

    def remove_grant(self):
        self.financial_params["grant_share"] = 0.0


# TODO if linear fit yields same results this can be deleted
def GoalSeek(fun, goal, x0, fTol=0.0001, MaxIter=1000):
    """
    code taken from https://github.com/DrTol/GoalSeek_Python
    Copyright (c) 2019 Hakan İbrahim Tol
    """
    # Goal Seek function of Excel
    #   via use of Line Search and Bisection Methods

    # Inputs
    #   fun     : Function to be evaluated
    #   goal    : Expected result/output
    #   x0      : Initial estimate/Starting point

    # Initial check
    if fun(x0) == goal:
        print("Exact solution found")
        return x0

    # Line Search Method
    step_sizes = np.logspace(-1, 4, 6)
    scopes = np.logspace(1, 5, 5)

    vFun = np.vectorize(fun)

    for scope in scopes:
        break_nested = False
        for step_size in step_sizes:
            cApos = np.linspace(x0, x0 + step_size * scope, int(scope))
            cAneg = np.linspace(x0, x0 - step_size * scope, int(scope))

            cA = np.concatenate((cAneg[::-1], cApos[1:]), axis=0)

            fA = vFun(cA) - goal

            if np.any(np.diff(np.sign(fA))):
                index_lb = np.nonzero(np.diff(np.sign(fA)))

                if len(index_lb[0]) == 1:
                    index_ub = index_lb + np.array([1])

                    x_lb = np.ndarray.item(np.array(cA)[index_lb])
                    x_ub = np.ndarray.item(np.array(cA)[index_ub])
                    break_nested = True
                    break
                else:  # Two or more roots possible
                    index_ub = index_lb + np.array([1])

                    print("Other solution possible at around, x0 = ", np.array(cA)[index_lb[0][1]])

                    x_lb = np.ndarray.item()(np.array(cA)[index_lb[0][0]])
                    x_ub = np.ndarray.item()(np.array(cA)[index_ub[0][0]])
                    break_nested = True
                    break

        if break_nested:
            break
    if not x_lb or not x_ub:
        print("No Solution Found")
        return

    # Bisection Method
    iter_num = 0
    error = 10

    while iter_num < MaxIter and fTol < error:
        x_m = (x_lb + x_ub) / 2
        f_m = fun(x_m) - goal

        error = abs(f_m)

        if (fun(x_lb) - goal) * (f_m) < 0:
            x_ub = x_m
        elif (fun(x_ub) - goal) * (f_m) < 0:
            x_lb = x_m
        elif f_m == 0:
            print("Exact spolution found")
            return x_m
        else:
            print("Failure in Bisection Method")

        iter_num += 1

    return x_m
